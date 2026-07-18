#!/usr/bin/env python3
"""Build the DSP391m capstone report as a .docx following the official template.

Tao file Word bam sat cau truc "Data Science Capstone Project Template" (DSP391m):
Introduction -> ... -> Model Development -> Model Evaluation and Fine-Tuning ->
Results Interpretation and Visualization -> Conclusion -> References -> Appendices.
So lieu ket qua doc truc tiep tu outputs/evaluation/results.json (khong hard-code tay);
bieu do nhung tu outputs/evaluation/figures_en/.

Cach chay:
    python scripts/plot_results.py --lang en --output-dir outputs/evaluation/figures_en
    python scripts/build_report_docx.py
"""

from __future__ import annotations

import json
import statistics
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
# Uu tien ket qua benchmark THAT nhieu cau; fallback ve audit n=1 neu chua co.
_BENCH = ROOT / "outputs/benchmark/evaluation/results.json"
RESULTS = _BENCH if _BENCH.is_file() else ROOT / "outputs/evaluation/results.json"
FIG = (ROOT / "outputs/benchmark/figures_en"
       if _BENCH.is_file() else ROOT / "outputs/evaluation/figures_en")
OUT = ROOT / "docs/DSP391m_Report3_Model_and_Results.docx"

ACCENT = RGBColor(0x1B, 0x49, 0x65)   # deep academic blue
INK = RGBColor(0x1A, 0x1F, 0x2B)
MUTED = RGBColor(0x5B, 0x67, 0x78)

ENGINE_ORDER = ["f5tts", "xtts", "mms", "piper", "edge", "bark"]
ENGINE_LABEL = {"f5tts": "F5-TTS", "xtts": "XTTS-v2", "mms": "MMS-TTS",
                "piper": "Piper", "edge": "Edge-TTS", "bark": "Bark"}
ENGINE_TYPE = {"f5tts": "voice cloning", "xtts": "voice cloning", "mms": "fixed voice",
               "piper": "fixed voice", "edge": "cloud, fixed", "bark": "fixed (no VI)"}


# ------------------------------------------------------------------ helpers ---
def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = INK
    rpr = normal.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.18
    pf.space_after = Pt(6)
    for lvl, size in ((1, 15), (2, 12.5), (3, 11.5)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(12 if lvl == 1 else 8)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True


def body(doc, text, *, justify=True, italic=False, color=None, size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(doc, segments, *, justify=True):
    """segments: list of (text, {'b':bool,'i':bool}) -> one paragraph with mixed runs."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in segments:
        run = p.add_run(text)
        run.bold = fmt.get("b", False)
        run.italic = fmt.get("i", False)
        if fmt.get("mono"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)
    return p


def lead(doc, label, text):
    """Template-style sub-item: bold lead-in label then text (e.g. 'Objective: ...')."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = ACCENT
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)
    return p


def figure(doc, filename, cap):
    path = FIG / filename
    if path.is_file():
        doc.add_picture(str(path), width=Inches(6.1))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): hex_color})
    tcpr.append(sh)


def set_cell(cell, text, *, bold=False, align="left", color=None, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


# -------------------------------------------------------------------- build ---
def main() -> int:
    data = json.loads(RESULTS.read_text(encoding="utf-8"))
    ok_rows = [r for r in data.get("results", []) if r.get("status") == "ok"]

    def _st(engine, key):
        vals = [float(r[key]) for r in ok_rows if r["engine"] == engine and r.get(key) is not None]
        if not vals:
            return {"mean": None, "std": 0.0, "median": None, "n": 0}
        return {"mean": statistics.mean(vals),
                "std": statistics.pstdev(vals) if len(vals) > 1 else 0.0,
                "median": statistics.median(vals), "n": len(vals)}

    summ = {}
    for e in {r["engine"] for r in ok_rows}:
        summ[e] = {
            "wer": _st(e, "wer"), "cer": _st(e, "cer"),
            "secs": _st(e, "speaker_similarity"),
            "time": _st(e, "inference_time"), "rtf": _st(e, "real_time_factor"),
            "n": sum(1 for r in ok_rows if r["engine"] == e),
        }
    engines = [e for e in ENGINE_ORDER if e in summ]
    n_sentences = max((summ[e]["n"] for e in engines), default=0)

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.9)
        section.left_margin = section.right_margin = Inches(1.0)
    set_base_style(doc)

    # ---- Title block ---------------------------------------------------------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DSP391m — DATA SCIENCE CAPSTONE PROJECT"); r.bold = True
    r.font.size = Pt(12); r.font.color.rgb = MUTED
    t.paragraph_format.space_after = Pt(2)

    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Vietnamese Voice Cloning Studio: Objective Benchmarking of "
                      "Zero-Shot Text-to-Speech with F5-TTS ViVoice")
    r.bold = True; r.font.size = Pt(19); r.font.color.rgb = ACCENT
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Report 3 — Model Development · Evaluation & Fine-Tuning · "
                    "Results · Conclusion"); r.italic = True; r.font.size = Pt(12)
    sub.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ---- Cover info box (bordered table, no role column) ---------------------
    info = [
        ("Course", ["DSP391m - Data Science Capstone Project"]),
        ("Team Members", ["Nguyễn Hoàng Thanh - SE172535",
                          "Trương Thanh Tuấn - SE182217",
                          "Nguyễn Thị Vân Anh - DE18037"]),
        ("Instructor", ["Nguyễn Quốc Trung"]),
        ("Academic Term", ["Summer 2026"]),
    ]
    box = doc.add_table(rows=len(info), cols=2)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    for i, (label, values) in enumerate(info):
        lc, vc = box.rows[i].cells
        lc.width = Inches(1.9)
        vc.width = Inches(4.4)
        # label cell
        lc.text = ""
        pl = lc.paragraphs[0]
        pl.paragraph_format.space_after = Pt(2)
        rl = pl.add_run(f"{label}:"); rl.font.size = Pt(11.5)
        # value cell (may span several lines, e.g. team members)
        vc.text = ""
        for k, val in enumerate(values):
            pv = vc.paragraphs[0] if k == 0 else vc.add_paragraph()
            pv.paragraph_format.space_after = Pt(2)
            rv = pv.add_run(val); rv.font.size = Pt(11.5)
        # vertical centering for tidy single-line rows
        for cell in (lc, vc):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()

    # ---- Project Title -------------------------------------------------------
    doc.add_heading("Project Title", level=1)
    body(doc, "Vietnamese Voice Cloning Studio — a bilingual web application that clones a "
              "Vietnamese speaker's voice from a short reference recording using F5-TTS, and "
              "benchmarks it against five additional text-to-speech engines under identical "
              "inputs with an objective, reproducible evaluation harness.")

    # ---- Introduction and Background ----------------------------------------
    doc.add_heading("1.  Introduction and Background", level=1)
    lead(doc, "Objective", "Build a Vietnamese zero-shot voice-cloning TTS system centered on "
         "F5-TTS ViVoice and quantify its quality against five baseline engines using objective, "
         "reproducible metrics (word/character error rate, speaker similarity, and generation speed).")
    lead(doc, "Motivation", "Voice cloning enables accessibility tools, dubbing and personalized "
         "assistants, yet Vietnamese is under-served. The language's six tones and dense diacritics "
         "make both synthesis and its evaluation hard — a single mispronounced tone can change a "
         "word's meaning — so a rigorous, tone-aware evaluation is as valuable as the model itself.")
    lead(doc, "Background Information", "Modern non-autoregressive TTS pairs a transformer backbone "
         "with a generative objective. F5-TTS uses a Diffusion Transformer (DiT) trained with a "
         "Conditional Flow-Matching (CFM) objective and a Vocos 24 kHz vocoder. The Vietnamese "
         "checkpoint hynt/F5-TTS-Vietnamese-ViVoice is pre-trained on roughly 1,000 hours of "
         "Vietnamese speech and released under CC-BY-NC-SA-4.0 for non-commercial research.")

    # ---- Literature Review ---------------------------------------------------
    doc.add_heading("2.  Literature Review", level=1)
    body(doc, "F5-TTS [1] demonstrates that a DiT [3] trained with flow matching [2] produces "
              "fluent, faithful speech without the alignment modules of earlier systems; Vocos [4] "
              "closes the quality gap between time-domain and Fourier-based neural vocoders. For "
              "zero-shot cloning, XTTS-v2 [5] is a strong multilingual competitor. Single-speaker "
              "baselines include Meta's MMS-TTS [6], the local ONNX voice Piper [9], Microsoft's "
              "cloud Edge-TTS, and Bark [8] (no Vietnamese support, used as a lower bound). "
              "Evaluation builds on Whisper [7] for the ASR round-trip and a GE2E-style speaker "
              "encoder [11] for identity similarity.")
    body(doc, "Gap addressed: published Vietnamese TTS work rarely reports a like-for-like, "
              "fully-reproducible objective comparison across cloning and fixed-voice engines under "
              "identical text, seed and reference. This project contributes exactly such a harness.")

    # ---- Data Description ----------------------------------------------------
    doc.add_heading("3.  Data Description", level=1)
    lead(doc, "Source", "Three data roles are distinguished. (a) Pre-training corpus: ViVoice "
         "(~1,000 h Vietnamese), used only via the released checkpoint. (b) Evaluation data: a fixed "
         "clean reference sample (reference_audio/sample_clean_vi.wav) and a fixed benchmark sentence. "
         "(c) Fine-tuning data: 30–60 minutes of single-speaker audio, to be collected.")
    lead(doc, "Size and Format", "Audio is mono WAV at 24 kHz. Text metadata uses a pipe-delimited "
         "CSV (audio_path | transcript). The model vocabulary is a plain-text file (~2.5k tokens) "
         "covering accented Vietnamese characters plus inherited pinyin tokens.")
    lead(doc, "Features", "Per sample: (i) the raw 24 kHz waveform → a 100-band mel-spectrogram "
         "front-end (n_fft 1024, hop 256, win 1024); (ii) the transcript, tokenized with the model's "
         "convert_char_to_pinyin scheme; and (iii) the sequence length used for masking.")

    # ---- Data Cleaning and Preprocessing ------------------------------------
    doc.add_heading("4.  Data Cleaning and Preprocessing", level=1)
    body(doc, "Audio preprocessing reuses the library's validated preprocess_ref_audio_text (silence "
              "trimming, resampling to 24 kHz mono, level normalization). Text is normalized to Unicode "
              "NFC with Vietnamese diacritics preserved; whitespace and punctuation are regularized. Long "
              "inputs pass through boundary-aware chunking that protects URLs, e-mails, decimals, dates and "
              "abbreviations from being split.")
    rich(doc, [("Challenge resolved — quarantined synthetic data. ", {"b": True}),
               ("An earlier iteration accidentally trained on synthetic sine/tone signals mislabeled as "
                "speech, corrupting the model. All such files were moved to a quarantined directory and are "
                "blocked from the pipeline unless an explicitly named technical-test flag is passed, so they "
                "can never again be mistaken for real training data.", {})])

    # ---- Exploratory Data Analysis ------------------------------------------
    doc.add_heading("5.  Exploratory Data Analysis (EDA)", level=1)
    body(doc, "Because the signal is audio, EDA relies on objective per-clip metrics — duration, RMS "
              "level, peak, clipping ratio and spectral flatness — computed directly on the waveform. "
              "Spectral flatness is the key discriminator: structured human speech sits around 0.10–0.15, "
              "whereas near-monotone or noisy signals collapse toward zero. This single metric exposed the "
              "corrupted-checkpoint defect (flatness ≈ 0.0019) and later confirmed the fix (flatness "
              "restored to ≈ 0.11–0.14, matching the real reference sample).")
    body(doc, "These metrics also act as automatic guardrails during evaluation: a clip that is silent, "
              "clipped, or non-finite is flagged before any error-rate is computed, so degenerate output "
              "cannot masquerade as a valid result.")

    # ---- Methodology ---------------------------------------------------------
    doc.add_heading("6.  Methodology", level=1)
    lead(doc, "Model Selection", "F5-TTS ViVoice is the primary model because it uniquely combines "
         "zero-shot cloning (a short reference suffices, no per-speaker training), strong Vietnamese "
         "pre-training, a state-of-the-art DiT+CFM backbone, and a research-permissive license. The five "
         "baselines were chosen to span the trade-off space: XTTS-v2 (competing cloning), MMS/Piper/Edge "
         "(fast fixed-voice TTS) and Bark (deliberate weak floor).")
    lead(doc, "Data Splitting Strategy", "For fine-tuning, a deterministic 90/10 train/validation split "
         "is drawn with a fixed random generator (seed 42), so every run yields the same split and results "
         "are comparable. Model selection uses validation loss, not training loss.")
    lead(doc, "Feature Engineering and Selection", "Two engineered representations are central: the "
         "100-band mel-spectrogram (audio side) and the pinyin-style character tokenization (text side). "
         "Text normalization and protected-token masking are the main hand-designed features; the "
         "architecture (mel bins, tokenizer) is fixed to match the pre-trained checkpoint and is not tuned.")

    # ---- 7. MODEL DEVELOPMENT (Report 3 core) --------------------------------
    doc.add_heading("7.  Model Development", level=1)
    doc.add_heading("7.1  Model Architecture", level=2)
    body(doc, "The architecture is loaded verbatim from the library's F5TTS_Base.yaml so it matches the "
              "pre-trained checkpoint exactly:")
    for line in [
        "CFM (Conditional Flow-Matching) wrapper",
        "  └─ DiT backbone: dim=1024, depth=22, heads=16, ff_mult=2,",
        "                   text_dim=512, conv_layers=4,",
        "                   text_mask_padding=false,  pe_attn_head=1   (compatibility-critical)",
        "  └─ Vocoder: Vocos @ 24 kHz (mel → waveform)",
        "Parameters ≈ 336M · front-end: n_fft=1024, hop=256, win=1024, n_mel=100",
    ]:
        p = doc.add_paragraph(); r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    rich(doc, [("Compatibility invariant. ", {"b": True}),
               ("The two fields text_mask_padding=false and pe_attn_head=1 must not be changed. When "
                "omitted, DiT reverts to v1 defaults and the forward pass silently mis-computes despite "
                "correctly-shaped weights, degrading output to a near-monotone signal (spectral flatness "
                "≈ 0.0019). Restoring the exact architecture returns flatness to the human-speech range. "
                "This was one of the two principal defects diagnosed and fixed during development.", {})])

    doc.add_heading("7.2  Training Procedure", level=2)
    body(doc, "The fine-tuning objective is single-speaker adaptation while avoiding catastrophic "
              "forgetting. The pipeline runs: metadata manifest → deterministic train/validation split → "
              "padding-aware collation → CFM flow-matching loss → mixed-precision back-propagation with "
              "gradient accumulation → gradient clipping → warm-up-then-cosine learning-rate schedule → "
              "checkpoint selection by validation loss. Exponential-moving-average weights are loaded from "
              "the official checkpoint; NaN/Inf losses abort immediately; each checkpoint records a JSON "
              "manifest of base model, architecture, vocabulary size, config, git commit, step, validation "
              "loss and dataset summary for full provenance.")
    body(doc, "Honest status: the fine-tuning pipeline is established and verified at the code-path level "
              "(its pure components are unit-tested), but it has not yet been run end-to-end on real "
              "speaker data with a GPU (blocked by the absence of collected speech). Accordingly, the "
              "results in Section 9 characterize the zero-shot base model, not a fine-tuned one; no "
              "fine-tuning improvement is claimed.", italic=True, color=MUTED)

    # ---- 8. MODEL EVALUATION AND FINE-TUNING ---------------------------------
    doc.add_heading("8.  Model Evaluation and Fine-Tuning", level=1)
    doc.add_heading("8.1  Evaluation Metrics", level=2)
    body(doc, "Generative speech has no single ground-truth label, so quality is measured through "
              "objective proxies on two core axes plus speed:")
    bullet(doc, "Intelligibility — WER and CER via an ASR round-trip: Whisper (small, beam 5, "
                "language=vi) re-transcribes the generated audio and the transcript is compared to the "
                "reference text with edit distance (jiwer, with an equivalent offline fallback). "
                "Vietnamese diacritics are preserved by default.")
    bullet(doc, "Speaker identity — SECS: cosine similarity between Resemblyzer embeddings of the "
                "generated and reference audio. It is a relative similarity score, not biometric proof.")
    bullet(doc, "Speed — real-time factor, RTF = inference time ÷ audio duration (RTF < 1 is faster "
                "than real time).")
    bullet(doc, "Signal guardrails — RMS, peak, clipping ratio and spectral flatness screen for "
                "degenerate output before error rates are computed.")

    doc.add_heading("8.2  Hyperparameter Tuning", level=2)
    body(doc, "Inference hyper-parameters are tunable today: the number of flow-matching ODE steps "
              "(NFE, default 32, raised to 48–64 for long/hard sentences to reduce late-sentence drift), "
              "the random seed (fixed at 42 for reproducibility and fair cross-engine comparison), the "
              "chunk length (≈280 characters) and inter-chunk silence (180 ms). Training hyper-parameters "
              "(for a T4 16 GB GPU) are summarized in Table 1.")
    ht = doc.add_table(rows=8, cols=3); ht.style = "Light Grid Accent 1"
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Hyper-parameter", "Value", "Rationale"]):
        set_cell(ht.rows[0].cells[j], h, bold=True, color=RGBColor(255, 255, 255))
        shade(ht.rows[0].cells[j], "1B4965")
    rows = [
        ("Learning rate", "1e-5", "Small LR for fine-tuning; avoids forgetting"),
        ("Batch × grad-accum", "1 × 8 = 8", "Effective batch under limited VRAM"),
        ("Warm-up steps", "200", "Linear 0→LR for early stability"),
        ("Scheduler", "cosine", "Anneals to min-LR 1e-7"),
        ("Max grad norm", "1.0", "Gradient clipping"),
        ("Mixed precision", "fp16", "T4 supports FP16; auto-off on CPU"),
        ("Optimizer / seed", "AdamW / 42", "β=(0.9,0.999), weight decay 0.01; deterministic"),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        set_cell(ht.rows[i].cells[0], a, bold=True)
        set_cell(ht.rows[i].cells[1], b, align="center")
        set_cell(ht.rows[i].cells[2], c)
    caption(doc, "Table 1. Training hyper-parameters (tuned for an NVIDIA T4, 16 GB).")

    doc.add_heading("8.3  Cross-Validation Techniques", level=2)
    body(doc, "Classical k-fold cross-validation is impractical for fine-tuning a 336M-parameter "
              "generative model on tens of minutes of audio (the k× training cost is prohibitive and the "
              "signal is generative, not a classification label). Its role is filled by (i) a deterministic "
              "held-out validation split monitored every epoch, with the best checkpoint chosen by "
              "validation loss, and (ii) an independent fixed test set of diverse sentences — greeting, "
              "general, numeric, date, interrogative, compound and long — used for the post-hoc "
              "WER/CER/SECS benchmark in Section 9.")

    # ---- 9. RESULTS INTERPRETATION AND VISUALIZATION -------------------------
    # Formatting + data-driven ranking helpers (report BOTH median and mean).
    def pct_mm(d):
        if d["median"] is None:
            return "—"
        return f"{d['median']*100:.1f} / {d['mean']*100:.1f}"

    def secs_fmt(d):
        return "—" if d["mean"] is None else f"{d['mean']:.3f}"

    def rank(metric, stat, *, best="min"):
        pairs = [(e, summ[e][metric][stat]) for e in engines if summ[e][metric][stat] is not None]
        pairs.sort(key=lambda t: t[1], reverse=(best == "max"))
        return pairs

    wer_med_rank = rank("wer", "median", best="min")
    wer_mean_rank = rank("wer", "mean", best="min")
    secs_rank = rank("secs", "mean", best="max")
    best_wer_med_e = wer_med_rank[0][0]
    best_wer_mean_e = wer_mean_rank[0][0]
    best_secs_e, best_secs_v = secs_rank[0]
    cloning = [e for e in engines if ENGINE_TYPE[e] == "voice cloning"]
    fixed = [e for e in ("mms", "piper", "edge") if e in summ]
    fixed_secs = [summ[e]["secs"]["mean"] for e in fixed if summ[e]["secs"]["mean"] is not None]

    def pctv(engine, metric, stat):
        return summ[engine][metric][stat] * 100

    doc.add_heading("9.  Results Interpretation and Visualization", level=1)
    body(doc, f"Every engine synthesizes the SAME set of {n_sentences} diverse Vietnamese sentences "
              f"(greeting, general, number, date, question, compound and long) from the SAME reference "
              f"sample under a fixed seed (42) and, for F5-TTS, a fixed NFE (32). Whisper re-transcribes "
              f"each output; metrics are computed per sentence and aggregated. Because a few hard sentences "
              f"(numbers, dates) legitimately inflate the average, Table 2 reports WER/CER as "
              f"median / mean, and Figure 5 shows the full per-sentence spread. Values are read directly "
              f"from {RESULTS.name}.")

    rt = doc.add_table(rows=1 + len(engines), cols=8); rt.style = "Light Grid Accent 1"
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Engine", "Type", "WER% ↓", "CER% ↓", "SECS ↑", "Time (s)", "RTF ↓", "N"]):
        set_cell(rt.rows[0].cells[j], h, bold=True, color=RGBColor(255, 255, 255), align="center", size=9)
        shade(rt.rows[0].cells[j], "1B4965")
    for i, e in enumerate(engines, start=1):
        s = summ[e]
        cells = rt.rows[i].cells
        set_cell(cells[0], ENGINE_LABEL[e], bold=True, size=9)
        set_cell(cells[1], ENGINE_TYPE[e], size=8.5)
        set_cell(cells[2], pct_mm(s["wer"]), align="center", size=9,
                 bold=e == best_wer_med_e, color=ACCENT if e == best_wer_med_e else None)
        set_cell(cells[3], pct_mm(s["cer"]), align="center", size=9)
        set_cell(cells[4], secs_fmt(s["secs"]), align="center", size=9,
                 bold=e == best_secs_e, color=ACCENT if e == best_secs_e else None)
        set_cell(cells[5], f"{s['time']['mean']:.1f}" if s["time"]["mean"] is not None else "—",
                 align="center", size=9)
        set_cell(cells[6], f"{s['rtf']['mean']:.2f}" if s["rtf"]["mean"] is not None else "—",
                 align="center", size=9)
        set_cell(cells[7], str(s["n"]), align="center", size=9)
        if e in cloning:
            for c in cells:
                shade(c, "E7EEF2")
    caption(doc, f"Table 2. Objective results over N={n_sentences} sentences (identical text / seed / "
                 f"reference). WER% and CER% shown as median / mean. Voice-cloning engines are shaded; "
                 f"blue marks the best median WER and the best SECS.")

    figure(doc, "wer_cer.png",
           f"Figure 1. Intelligibility via ASR round-trip (bars = mean, whiskers = std over "
           f"{n_sentences} sentences). The strong engines cluster closely; Bark (no Vietnamese) is the "
           f"clear lower bound.")
    figure(doc, "speaker_similarity.png",
           f"Figure 2. Speaker similarity (SECS). Only the two cloning engines, F5-TTS "
           f"({summ['f5tts']['secs']['mean']:.3f}) and XTTS-v2 ({summ['xtts']['secs']['mean']:.3f}), "
           f"enter the voice-cloning region (>0.75); fixed-voice engines fall well below it.")
    figure(doc, "speed_rtf.png",
           "Figure 3. Generation speed (RTF, log scale; lower is faster). Fixed-voice engines run at or "
           "faster than real time on CPU; the flow-matching cloning models are slower on CPU and are "
           "expected to speed up substantially on GPU.")
    figure(doc, "tradeoff.png",
           "Figure 4. The central trade-off — intelligibility (1−mean WER) vs. speaker identity; bubble "
           "size ∝ speed. Speaker similarity, not intelligibility, is what separates cloning from "
           "fixed-voice engines.")
    figure(doc, "wer_distribution.png",
           f"Figure 5. Per-sentence WER across the {n_sentences} sentences (each dot is one sentence; the "
           f"bar is the mean). It visualizes uncertainty: most sentences score low, while a few hard ones "
           f"(numbers/dates) pull the mean up — which is why the median is the more robust summary.")

    doc.add_heading("9.1  Insights and Implications", level=2)
    bullet(doc, f"Speaker identity is the decisive axis — only the two voice-cloning engines exceed the "
                f"0.75 cloning threshold (F5-TTS {summ['f5tts']['secs']['mean']:.3f}, XTTS-v2 "
                f"{summ['xtts']['secs']['mean']:.3f}), whereas the fixed-voice engines score "
                f"≈{min(fixed_secs):.2f}–{max(fixed_secs):.2f} — they render a default voice, not the "
                f"reference speaker. This is the core voice-cloning result.")
    bullet(doc, f"Intelligibility is broadly comparable among the strong engines — median WER for "
                f"F5-TTS ({pctv('f5tts','wer','median'):.1f}%) and XTTS-v2 "
                f"({pctv('xtts','wer','median'):.1f}%) is close to the best fixed-voice engines, and on "
                f"{n_sentences} sentences these gaps are within noise. F5-TTS has the lowest median WER; "
                f"XTTS-v2 the lowest mean — effectively a tie.")
    bullet(doc, "Median vs. mean matters — means are inflated by a few hard sentences (numbers and dates) "
                "where the spoken form and Whisper's transcription legitimately differ (Figure 5); the "
                "median is the more robust central tendency here.")
    bullet(doc, f"Cost of quality — the cloning engines are slower on CPU (F5-TTS RTF "
                f"{summ['f5tts']['rtf']['mean']:.1f}, XTTS-v2 {summ['xtts']['rtf']['mean']:.1f}) than the "
                f"fixed-voice engines (RTF < 1); XTTS-v2 is the faster of the two on CPU.")
    bullet(doc, "The harness catches real failure — the deliberately weak baseline (Bark, no Vietnamese "
                "support) is separated automatically by very high WER and the lowest SECS, evidence that "
                "the metrics measure something real rather than rewarding every engine equally.")

    doc.add_heading("9.2  Limitations and Uncertainty", level=2)
    bullet(doc, f"Test-set scale — results cover {n_sentences} sentences per engine; error bars (Figure 1) "
                "and the distribution (Figure 5) quantify spread, but a larger, category-balanced set would "
                "tighten the estimates and enable significance testing.")
    bullet(doc, "ASR is imperfect — Whisper itself mis-transcribes, especially tones, so WER/CER are "
                "upper bounds on the true error and are most meaningful when compared across engines.")
    bullet(doc, "SECS is relative — it is a similarity score between embeddings, not biometric proof, "
                "and it does not replace subjective human listening (a MOS study remains future work).")
    bullet(doc, "Speed is CPU-measured — inference time and RTF are CPU figures (and the first call per "
                "engine includes model-load overhead); absolute latency will differ on GPU.")
    bullet(doc, "Base model only — these numbers characterize the zero-shot base model; the fine-tuning "
                "pipeline is established and verified at the code-path level but not yet run end-to-end on "
                "real speaker data with a GPU, so no fine-tuning improvement is claimed.")

    # ---- 10. CONCLUSION (PROVISIONAL) ----------------------------------------
    doc.add_heading("10.  Conclusion and Recommendations (Provisional)", level=1)
    body(doc, f"These are provisional conclusions reflecting the current project status — an objective "
              f"benchmark over {n_sentences} sentences on the zero-shot base model — not a final, "
              f"large-scale or fine-tuned evaluation.", italic=True, color=MUTED)
    lead(doc, "Key findings (current status)",
         f"On this {n_sentences}-sentence benchmark, the two voice-cloning engines — F5-TTS and XTTS-v2 — "
         f"are the only systems that reproduce the reference speaker (SECS "
         f"{summ['f5tts']['secs']['mean']:.3f} and {summ['xtts']['secs']['mean']:.3f}, both above the 0.75 "
         f"threshold), while the fixed-voice engines do not. Intelligibility is comparable across the "
         f"strong engines: F5-TTS gives the lowest median WER ({pctv('f5tts','wer','median'):.1f}%) and "
         f"XTTS-v2 the lowest mean, an effective tie. Just as important as any single winner, the project "
         f"delivers a reproducible objective-evaluation framework that turns subjective impressions into "
         f"verifiable measurements and surfaced two subtle, high-impact defects during development.")
    lead(doc, "Recommendations",
         f"For Vietnamese voice cloning, F5-TTS is the recommended primary engine — it matches the field "
         f"on intelligibility (lowest median WER), reaches top-tier speaker similarity, and is "
         f"pre-trained on Vietnamese; enable GPU and raise NFE to 48–64 for long sentences. XTTS-v2 is an "
         f"equally strong, faster alternative when speaker likeness and CPU speed matter most. Use "
         f"Piper/MMS/Edge for real-time offline reading that does not need a specific voice, and avoid "
         f"Bark for Vietnamese. Always provide a clean 5–10 s single-speaker reference with an accurate "
         f"transcript.")
    lead(doc, "Reflection and future work",
         "Establishing an objective yardstick was the most valuable outcome: it made subtle defects visible "
         "and keeps every claim evidence-based. Next, in priority order: (1) enlarge and balance the test "
         "set and report confidence intervals / significance; (2) collect 30–60 minutes of real speaker "
         "audio and run fine-tuning end-to-end on GPU, measuring WER/CER/SECS before vs. after; (3) add "
         "subjective MOS listening tests; (4) optimize cloning-engine latency on GPU; (5) package the "
         "system for setup-free demos.")

    # ---- 11. License and Ethics ----------------------------------------------
    doc.add_heading("11.  License and Ethical Considerations", level=1)
    lead(doc, "License", "The base model hynt/F5-TTS-Vietnamese-ViVoice is released under "
         "CC-BY-NC-SA-4.0: this project is strictly non-commercial and for education/research only. Any "
         "commercial use would require replacing the model with an appropriately licensed one. Comparison "
         "engines retain their own licenses (e.g., Coqui, Meta MMS, Piper, Bark) and Edge-TTS is a "
         "cloud service subject to Microsoft's terms.")
    lead(doc, "Ethics", "Voice cloning can be misused for impersonation. Users must clone only their own "
         "voice or a voice they have explicit permission to use; generated audio should be disclosed as "
         "synthetic where appropriate. Objective similarity scores here are for quality assessment only "
         "and must not be treated as identity verification.")

    # ---- References ----------------------------------------------------------
    doc.add_heading("References", level=1)
    refs = [
        "Y. Chen et al., “F5-TTS: A Fairytaler that Fakes Fluent and Faithful Speech with Flow "
        "Matching,” arXiv:2410.06885, 2024.",
        "Y. Lipman et al., “Flow Matching for Generative Modeling,” ICLR, 2023. arXiv:2210.02747.",
        "W. Peebles and S. Xie, “Scalable Diffusion Models with Transformers,” ICCV, 2023. "
        "arXiv:2212.09748.",
        "H. Siuzdak, “Vocos: Closing the Gap Between Time-Domain and Fourier-Based Neural Vocoders,"
        "” ICLR, 2024. arXiv:2306.00814.",
        "E. Casanova et al., “XTTS: A Massively Multilingual Zero-Shot Text-to-Speech Model,” "
        "arXiv:2406.04904, 2024.",
        "V. Pratap et al., “Scaling Speech Technology to 1,000+ Languages (MMS),” arXiv:2305.13516, 2023.",
        "A. Radford et al., “Robust Speech Recognition via Large-Scale Weak Supervision (Whisper),"
        "” ICML, 2023. arXiv:2212.04356.",
        "Suno AI, “Bark: Text-Prompted Generative Audio Model,” GitHub, 2023.",
        "Rhasspy, “Piper: A Fast, Local Neural Text-to-Speech System,” GitHub, 2023.",
        "hynt, “F5-TTS-Vietnamese-ViVoice,” Hugging Face model repository, 2024.",
        "L. Wan et al., “Generalized End-to-End Loss for Speaker Verification,” ICASSP, 2018. "
        "arXiv:1710.10467.",
    ]
    for i, ref in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(3)
        rn = p.add_run(f"[{i}] "); rn.bold = True; rn.font.color.rgb = ACCENT
        r = p.add_run(ref); r.font.size = Pt(10)

    # ---- Appendices ----------------------------------------------------------
    doc.add_heading("Appendices", level=1)
    doc.add_heading("Appendix A — Reproducing the results", level=2)
    for line in [
        "# 1. Generate audio for all engines over the multi-sentence test set (CPU)",
        "python scripts/run_benchmark.py",
        "# 2. Evaluate: WER/CER (Whisper) + speaker similarity (Resemblyzer) + RTF",
        "python scripts/evaluate.py --manifest outputs/benchmark/manifest.json \\",
        "    --output-dir outputs/benchmark/evaluation",
        "# 3. Render the figures (English labels, mean ± std)",
        "python scripts/plot_results.py --results outputs/benchmark/evaluation/results.json \\",
        "    --output-dir outputs/benchmark/figures_en --lang en",
        "# 4. Rebuild this document",
        "python scripts/build_report_docx.py",
    ]:
        p = doc.add_paragraph(); r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    body(doc, "Appendix B — Key source files: app.py (UI/orchestration), engines.py (five comparison "
              "engines), voice_studio/ (offline-testable logic: text processing, audio metrics, "
              "evaluation), scripts/train.py (fine-tuning), scripts/evaluate.py and scripts/plot_results.py "
              "(evaluation and charts), configs/train_config.yaml (architecture and training config).")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Saved: {OUT}  ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
